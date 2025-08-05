import requests
import subprocess
import logging
import tempfile
import os
import glob
import pandas as pd
import sqlite3
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta

# List of raw GitHub URLs to your scripts
GITHUB_SCRIPTS = [
    "https://raw.githubusercontent.com/MariaKlap/RI/main/EMAnews2.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/ECnews11.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/ICR.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/ICHnews.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/IS1.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/SWISS5.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/AT.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/GMP.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/EC-Updates.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/EC-Medical.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/FDAnews.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/RQAnews4.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/Topra.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/raps-2.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/WHOnews.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/CBGnewsfinal5win.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/main/HMA6news.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/BEnews1.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/CY.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/DE.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/DK3newswin.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/FInew.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/IE.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/Infarmed6news.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/Luxnews.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/MHRA.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/MHRANews.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/MHRAPolicy.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/Maltanews.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/Norwnews%20(2).py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/SEn.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/SEns.py",
    "https://raw.githubusercontent.com/MariaKlap/RI/refs/heads/main/SEnsa.py",
]

# Set up logging
log_file = os.path.join(os.getcwd(), "batch_run_log.txt")
logging.basicConfig(filename=log_file, level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')

def download_and_run_script(url):
    try:
        logging.info(f"📥 Downloading script: {url}")
        response = requests.get(url)
        response.raise_for_status()

        with tempfile.NamedTemporaryFile(delete=False, suffix=".py") as temp_file:
            temp_file.write(response.content)
            temp_file_path = temp_file.name

        logging.info(f"🚀 Running script: {url}")
        subprocess.run(["python", temp_file_path], check=True)
        logging.info(f"✅ Completed: {url}")
        
        os.remove(temp_file_path)
    except requests.RequestException as e:
        logging.error(f"❌ Download failed for {url}: {e}")
    except subprocess.CalledProcessError as e:
        logging.error(f"❌ Execution failed for {url}: {e}")

def combine_excel_files():
    """Combine all Excel files in the directory into one RI.xlsx, preserving None values"""
    combined_df = pd.DataFrame()
    excel_files = []
    
    try:
        logging.info("🔍 Searching for Excel files to combine...")
        excel_files = glob.glob(os.path.join(os.getcwd(), '*.xlsx'))

        # Filter out our output file if it exists
        excel_files = [f for f in excel_files if not f.endswith('RI.xlsx')]

        if not excel_files:
            logging.warning("⚠️ No Excel files found to combine")
            return False

        logging.info(f"📂 Found {len(excel_files)} Excel files to combine")

        for file in excel_files:
            try:
                df = pd.read_excel(file, keep_default_na=True)
                df['Source_File'] = os.path.basename(file)
                combined_df = pd.concat([combined_df, df], ignore_index=True)
                logging.info(f"➕ Added {file} to combined dataframe")
            except Exception as e:
                logging.error(f"❌ Error reading {file}: {e}")
                continue  # Continue with next file even if one fails

        # Enhanced Date Handling
        if 'Date' in combined_df.columns:
            try:
                # Step 1: Replace common placeholders with NA
                combined_df['Date'] = combined_df['Date'].replace(['None', 'N/A', 'NA', ''], pd.NA)
                
                # Step 2: Standardize separators (dots and dashes to slashes)
                combined_df['Date'] = combined_df['Date'].astype(str).str.replace('[.-]', '/', regex=True)
                
                # Step 3: Convert to datetime
                parsed_dates = pd.to_datetime(
                    combined_df['Date'],
                    errors='coerce',
                    dayfirst=True
                )
                
                combined_df['Date'] = parsed_dates
                
                # Step 4: Filter by date (keep last 12 months or None)
                one_year_ago = pd.Timestamp.now() - pd.DateOffset(months=12)
                initial_count = len(combined_df)
                
                combined_df = combined_df[
                    (combined_df['Date'].isna()) | 
                    (combined_df['Date'] >= one_year_ago)
                ]
                
                filtered_count = initial_count - len(combined_df)
                logging.info(f"🧹 Filtered out {filtered_count} records older than {one_year_ago.date()}")
                
                # Step 5: Format for output (preserve NaT as None)
                combined_df['Date'] = combined_df['Date'].dt.strftime('%d-%m-%Y').where(
                    combined_df['Date'].notna(), None
                )
                
                logging.info(f"✅ Processed 'Date' column. Kept {len(combined_df)} records.")
            except Exception as e:
                logging.error(f"❌ Failed to process 'Date' column: {e}")
                # Keep original dates if processing fails
                combined_df['Date'] = combined_df['Date'].astype(str)
        else:
            logging.warning("⚠️ 'Date' column not found in combined DataFrame. Skipping date filtering.")

        if not combined_df.empty:
            output_path = os.path.join(os.getcwd(), 'RI.xlsx')
            combined_df.to_excel(output_path, index=False, na_rep='None')
            logging.info(f"💾 Saved combined Excel to {output_path}")
            return True
        else:
            logging.warning("⚠️ No data to save - combined dataframe is empty")
            return False

    except Exception as e:
        logging.error(f"❌ Unexpected error in combine_excel_files: {e}")
        return False
        
    finally:
        # Cleanup or resource release could go here if needed
        logging.info("🏁 Finished combine_excel_files operation")

def convert_excel_to_db():
    """Convert RI.xlsx to RI.db SQLite database and RI.csv file, preserving None values"""
    try:
        excel_path = os.path.join(os.getcwd(), 'RI.xlsx')
        db_path = os.path.join(os.getcwd(), 'RI.db')
        csv_path = os.path.join(os.getcwd(), 'RI.csv')
        
        if not os.path.exists(excel_path):
            logging.warning("⚠️ RI.xlsx not found - cannot create database or CSV")
            return False
        
        logging.info("📊 Reading combined Excel file...")
        # Read Excel keeping None values
        df = pd.read_excel(excel_path, keep_default_na=True)
        
        if df.empty:
            logging.warning("⚠️ Excel file is empty - cannot create database or CSV")
            return False
        
        # Save as CSV
        logging.info("💾 Creating CSV file...")
        df.to_csv(csv_path, index=False, na_rep='None')
        logging.info(f"✅ Successfully created CSV file at {csv_path}")

        # Save as SQLite DB
        logging.info("💾 Creating SQLite database...")
        conn = sqlite3.connect(db_path)
        df = df.where(pd.notnull(df), None)  # Convert None to NULL
        df.to_sql('regulatory_intelligence', conn, if_exists='replace', index=False)
        conn.close()
        logging.info(f"✅ Successfully created database at {db_path}")

        return True

    except Exception as e:
        logging.error(f"❌ Error creating database or CSV: {e}")
        return False


def compare_with_github_csv():
    """
    Compare local RI.csv with GitHub RI.csv based on 'Article URL'.
    Write unmatched articles to News.xlsx.
    """
    try:
        local_csv_path = os.path.join(os.getcwd(), 'RI.csv')
        github_csv_url = "https://raw.githubusercontent.com/MariaKlap/Master-Script/refs/heads/main/RI.csv"
        output_excel_path = os.path.join(os.getcwd(), 'News.xlsx')

        if not os.path.exists(local_csv_path):
            logging.warning("⚠️ Local RI.csv not found. Skipping comparison.")
            return False

        # Load local CSV (assumed to be comma-separated)
        df_local = pd.read_csv(local_csv_path, keep_default_na=True)

        # Load GitHub CSV (semicolon-separated)
        df_github = pd.read_csv(
            github_csv_url,
            sep=';',
            keep_default_na=True,
            on_bad_lines='skip',  # if using Python < 3.10, replace with error_bad_lines=False
            quoting=1             # handle quoted fields properly
        )

        # Ensure 'Article URL' column exists
        if 'Article URL' not in df_local.columns or 'Article URL' not in df_github.columns:
            logging.error("❌ 'Article URL' column missing in one of the files.")
            return False

        # Compare: find rows in local not in GitHub
        unmatched_df = df_local[~df_local['Article URL'].isin(df_github['Article URL'])]

        if unmatched_df.empty:
            logging.info("✅ No unmatched articles found. News.xlsx will not be created.")
            return True

        # Ensure required columns exist
        required_columns = ['Title', 'Summary', 'Date', 'Article URL', 'Source_File']
        for col in required_columns:
            if col not in unmatched_df.columns:
                unmatched_df[col] = None

        # Select and save
        result_df = unmatched_df[required_columns]
        result_df.to_excel(output_excel_path, index=False)
        logging.info(f"📝 Unmatched articles saved to {output_excel_path}")
        return True

    except Exception as e:
        logging.error(f"❌ Error during comparison with GitHub RI.csv: {e}")
        return False

def export_news_to_docx():
    """Convert News.xlsx into a nicely formatted RI_News.docx with clickable links"""
    try:
        excel_path = os.path.join(os.getcwd(), 'News.xlsx')
        docx_path = os.path.join(os.getcwd(), 'RI_News.docx')

        if not os.path.exists(excel_path):
            logging.warning("⚠️ News.xlsx not found. Skipping DOCX export.")
            return False

        df = pd.read_excel(excel_path)

        if df.empty:
            logging.info("⚠️ News.xlsx is empty. No DOCX generated.")
            return False

        doc = Document()
        doc.add_heading("RI News", level=1)

        for _, row in df.iterrows():
            title = str(row.get("Title", "")).strip()
            summary = str(row.get("Summary", "")).strip()
            date = str(row.get("Date", "")).strip()
            url = str(row.get("Article URL", "")).strip()

            # Title with hyperlink
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("Title: ")
            run.bold = True

            # Create hyperlink
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), doc.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            new_run.append(rPr)

            text_elem = OxmlElement('w:t')
            text_elem.text = title
            new_run.append(text_elem)
            hyperlink.append(new_run)

            paragraph._p.append(hyperlink)

            # Summary
            doc.add_paragraph(f"Summary: {summary}")
            doc.add_paragraph(f"Date: {date}")
            doc.add_paragraph()  # empty line

        doc.save(docx_path)
        logging.info(f"📝 Exported news to {docx_path}")
        return True

    except Exception as e:
        logging.error(f"❌ Error exporting to DOCX: {e}")
        return False


def main():
    logging.info("=== Batch GitHub Execution Started ===")
    
    for url in GITHUB_SCRIPTS:
        download_and_run_script(url)
    
    logging.info("=== Combining Excel Files ===")
    if combine_excel_files():
        logging.info("=== Converting to Database and CSV ===")
        success = convert_excel_to_db()
        
        if success:
            # Ensure the CSV has been written and is ready
            csv_path = os.path.join(os.getcwd(), 'RI.csv')
            wait_time = 2  # seconds
            for _ in range(5):  # Retry up to 5 times
                if os.path.exists(csv_path):
                    logging.info(f"⏳ Waiting {wait_time} seconds to ensure RI.csv is ready...")
                    time.sleep(wait_time)
                    break
                else:
                    logging.warning(f"⏳ RI.csv not found yet. Retrying...")
                    time.sleep(1)
            
            logging.info("=== Comparing with GitHub RI.csv ===")
            if compare_with_github_csv():
                logging.info("=== Exporting unmatched news to DOCX ===")
                export_news_to_docx()
    
    logging.info("=== Batch GitHub Execution Completed ===")
    print("✅ All scripts from GitHub have been executed. Check 'batch_run_log.txt' for details.")



if __name__ == "__main__":
    main()
